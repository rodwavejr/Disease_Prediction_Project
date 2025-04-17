"""
Create a Word document from the project explanation markdown file.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_word(md_file, word_file):
    """
    Convert a markdown file to a Word document with formatting.
    
    Parameters:
    -----------
    md_file : str
        Path to the markdown file
    word_file : str
        Path to save the Word document
    """
    # Read markdown file
    with open(md_file, 'r') as f:
        md_content = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title page
    title = doc.add_heading('COVID-19 Detection from Unstructured Medical Text', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project Explanation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph()  # Add some space
    date = doc.add_paragraph('April 2025')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page break
    doc.add_page_break()
    
    # Process markdown content
    in_code_block = False
    current_list_level = 0
    
    for line in md_content:
        line = line.rstrip()
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], 1)
            
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
            
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
            
        # Handle code blocks
        elif line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Start a new code block
                current_code_block = []
            else:
                # End the code block and add it to the document
                code_para = doc.add_paragraph()
                code_text = code_para.add_run('\n'.join(current_code_block))
                code_text.font.name = 'Courier New'
                code_text.font.size = Pt(9)
                
        elif in_code_block:
            current_code_block.append(line)
            
        # Handle bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
        # Handle numbered lists
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            num, text = line.split('. ', 1)
            p = doc.add_paragraph(text, style='List Number')
            
        # Handle blank lines
        elif not line.strip():
            doc.add_paragraph()
            
        # Regular paragraph
        else:
            doc.add_paragraph(line)
    
    # Save the document
    doc.save(word_file)
    print(f"Word document saved to {word_file}")

if __name__ == "__main__":
    os.makedirs("presentations", exist_ok=True)
    
    md_file = "docs/project_explanation.md"
    word_file = "presentations/COVID19_Detection_Project.docx"
    
    markdown_to_word(md_file, word_file)