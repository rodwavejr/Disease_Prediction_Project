#!/usr/bin/env python3
"""
Convert Markdown file to PDF and DOCX formats
"""

import os
import markdown
import subprocess
from pathlib import Path

# File paths
md_file = '/Users/Apexr/Documents/Disease_Prediction_Project/COVID19_Project_Report_4page.md'
html_file = '/Users/Apexr/Documents/Disease_Prediction_Project/COVID19_Project_Report_4page.html'
pdf_file = '/Users/Apexr/Documents/Disease_Prediction_Project/COVID19_Project_Report_4page.pdf'
docx_file = '/Users/Apexr/Documents/Disease_Prediction_Project/COVID19_Project_Report_4page.docx'

def markdown_to_html(md_path, html_path):
    """Convert markdown to HTML with proper styling"""
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'codehilite']
    )
    
    # Add CSS styling
    styled_html = f"""<!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                line-height: 1.6;
                padding: 20px;
                max-width: 800px;
                margin: 0 auto;
                color: #333;
            }}
            h1, h2, h3, h4, h5, h6 {{
                margin-top: 24px;
                margin-bottom: 16px;
                font-weight: 600;
                line-height: 1.25;
            }}
            h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
            h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
            h3 {{ font-size: 1.25em; }}
            h4 {{ font-size: 1em; }}
            h5 {{ font-size: 0.875em; }}
            h6 {{ font-size: 0.85em; color: #6a737d; }}
            
            /* Enhanced code formatting */
            code, pre {{
                font-family: SFMono-Regular, Consolas, "Liberation Mono", Menlo, monospace;
            }}
            
            pre {{
                background-color: #f6f8fa;
                border-radius: 6px;
                padding: 16px;
                overflow: auto;
                line-height: 1.45;
                margin-bottom: 16px;
                border: 1px solid #e1e4e8;
            }}
            
            code {{
                background-color: #f6f8fa;
                padding: 0.2em 0.4em;
                margin: 0;
                font-size: 85%;
                border-radius: 3px;
            }}
            
            pre code {{
                background-color: transparent;
                padding: 0;
                margin: 0;
                font-size: 100%;
                word-break: normal;
                white-space: pre;
                border: none;
            }}
            
            /* Method section special formatting */
            .method-step {{
                margin-bottom: 20px;
                padding-left: 10px;
                border-left: 3px solid #0366d6;
            }}
            
            .method-step h4 {{
                color: #0366d6;
                margin-top: 0;
            }}
            
            .code-example {{
                display: flex;
                margin-bottom: 20px;
            }}
            
            .code-example-before, .code-example-after {{
                flex: 1;
                padding: 10px;
                border-radius: 6px;
                background-color: #f6f8fa;
                border: 1px solid #e1e4e8;
                margin: 5px;
            }}
            
            .code-example-before h4, .code-example-after h4 {{
                margin-top: 0;
                color: #0366d6;
            }}
            
            blockquote {{
                padding: 0 1em;
                color: #6a737d;
                border-left: 0.25em solid #dfe2e5;
                margin: 0 0 16px 0;
            }}
            
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-bottom: 16px;
            }}
            
            table th, table td {{
                padding: 8px 13px;
                border: 1px solid #dfe2e5;
            }}
            
            table th {{
                background-color: #f1f8ff;
                font-weight: 600;
            }}
            
            table tr {{
                background-color: #fff;
                border-top: 1px solid #c6cbd1;
            }}
            
            table tr:nth-child(2n) {{
                background-color: #f6f8fa;
            }}
            
            img {{
                max-width: 100%;
                box-sizing: border-box;
                display: block;
                margin: 20px auto;
                border-radius: 6px;
            }}
            
            hr {{
                height: 0.25em;
                padding: 0;
                margin: 24px 0;
                background-color: #e1e4e8;
                border: 0;
            }}
        </style>
    </head>
    <body>
        {html}
    </body>
    </html>"""
    
    # Write the HTML file
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(styled_html)
    
    print(f"HTML file created: {html_path}")

def html_to_pdf(html_path, pdf_path):
    """Convert HTML to PDF using wkhtmltopdf if available, otherwise try other methods"""
    try:
        # Try using wkhtmltopdf if available
        subprocess.run([
            'wkhtmltopdf',
            '--enable-local-file-access',
            '--page-size', 'Letter',
            '--margin-top', '20',
            '--margin-right', '20',
            '--margin-bottom', '20',
            '--margin-left', '20',
            html_path, 
            pdf_path
        ], check=True)
        print(f"PDF file created: {pdf_path}")
        return True
    except (subprocess.SubprocessError, FileNotFoundError):
        try:
            # Try using Chrome/Chromium if available (headless mode)
            subprocess.run([
                'chrome',
                '--headless',
                '--disable-gpu',
                f'--print-to-pdf={pdf_path}',
                html_path
            ], check=True)
            print(f"PDF file created: {pdf_path}")
            return True
        except (subprocess.SubprocessError, FileNotFoundError):
            # If neither method works, provide instructions
            print("Could not create PDF automatically. Please:")
            print(f"1. Open the HTML file ({html_path}) in a web browser")
            print("2. Use the browser's Print function (Cmd+P or Ctrl+P)")
            print("3. Choose 'Save as PDF' as the destination")
            print(f"4. Save to: {pdf_path}")
            return False

def html_to_docx(html_path, docx_path):
    """Convert HTML to DOCX using pandoc if available, otherwise provide instructions"""
    try:
        # Try using pandoc if available
        subprocess.run([
            'pandoc',
            html_path,
            '-o', docx_path,
            '-f', 'html',
            '-t', 'docx'
        ], check=True)
        print(f"Word document created: {docx_path}")
        return True
    except (subprocess.SubprocessError, FileNotFoundError):
        try:
            # Try alternative method: convert HTML to DOCX using python-docx
            from bs4 import BeautifulSoup
            from docx import Document
            from docx.shared import Pt, Inches
            
            # Parse HTML
            with open(html_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            # Create document
            doc = Document()
            
            # Process elements
            for elem in soup.body.children:
                if elem.name == 'h1':
                    doc.add_heading(elem.text, level=1)
                elif elem.name == 'h2':
                    doc.add_heading(elem.text, level=2)
                elif elem.name == 'h3':
                    doc.add_heading(elem.text, level=3)
                elif elem.name == 'p':
                    doc.add_paragraph(elem.text)
                # Add more element types as needed
            
            # Save document
            doc.save(docx_path)
            print(f"Word document created: {docx_path}")
            return True
        except (ImportError, Exception) as e:
            # If neither method works, provide instructions
            print(f"Could not create DOCX automatically: {e}")
            print("Please:")
            print(f"1. Open the HTML file ({html_path}) in a web browser")
            print("2. Copy all content (Cmd+A then Cmd+C)")
            print("3. Paste into a new Word document")
            print(f"4. Save as: {docx_path}")
            return False

if __name__ == "__main__":
    # Convert MD to HTML
    markdown_to_html(md_file, html_file)
    
    # Convert HTML to PDF
    html_to_pdf(html_file, pdf_file)
    
    # Convert HTML to DOCX
    html_to_docx(html_file, docx_file)
    
    print("\nConversion complete!")
    print(f"Files created in: {Path(md_file).parent}")